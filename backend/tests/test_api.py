import io
import json
import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
import pypdfium2 as pdfium
import docx

from database import Base, get_db
import models
import auth as auth_module
from main import app
from services.ai.base import AiProvider
import services.ai.router as ai_router
from services.ai.groq_provider import GroqProvider
from services.ai.openai_provider import OpenAiProvider

# Test SQLite DB in memory or temporary file
SQLALCHEMY_TEST_DATABASE_URL = "sqlite:///./test_app.db"
test_engine = create_engine(
    SQLALCHEMY_TEST_DATABASE_URL, connect_args={"check_same_thread": False}
)
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=test_engine)


def override_get_db():
    db = TestingSessionLocal()
    try:
        yield db
    finally:
        db.close()


app.dependency_overrides[get_db] = override_get_db


@pytest.fixture(autouse=True)
def setup_db():
    Base.metadata.drop_all(bind=test_engine)
    Base.metadata.create_all(bind=test_engine)
    yield
    Base.metadata.drop_all(bind=test_engine)


def _login_client(username: str, password: str) -> TestClient:
    test_client = TestClient(app)
    res = test_client.post("/api/auth/login", json={"username": username, "password": password})
    assert res.status_code == 200, res.text
    body = res.json()
    test_client.headers["Authorization"] = f"Bearer {body['access_token']}"
    test_client.test_username = username
    return test_client


@pytest.fixture
def client():
    db = TestingSessionLocal()
    try:
        auth_module.ensure_bootstrap_admin(db)
    finally:
        db.close()

    test_client = _login_client(auth_module.BOOTSTRAP_ADMIN_USERNAME, auth_module.BOOTSTRAP_ADMIN_PASSWORD)

    db = TestingSessionLocal()
    try:
        user = db.query(models.User).filter(models.User.username == auth_module.BOOTSTRAP_ADMIN_USERNAME).first()
        test_client.test_user_id = user.id
    finally:
        db.close()

    return test_client


def test_root(client):
    response = client.get("/")
    assert response.status_code == 200
    assert response.json()["status"] == "healthy"


def test_login_requires_correct_credentials():
    test_client = TestClient(app)
    db = TestingSessionLocal()
    try:
        auth_module.ensure_bootstrap_admin(db)
    finally:
        db.close()

    bad_res = test_client.post(
        "/api/auth/login",
        json={"username": auth_module.BOOTSTRAP_ADMIN_USERNAME, "password": "wrong-password"},
    )
    assert bad_res.status_code == 401

    good_res = test_client.post(
        "/api/auth/login",
        json={"username": auth_module.BOOTSTRAP_ADMIN_USERNAME, "password": auth_module.BOOTSTRAP_ADMIN_PASSWORD},
    )
    assert good_res.status_code == 200
    assert good_res.json()["is_admin"] is True

    # Protected routes reject requests with no token, and with a garbage token
    assert test_client.get("/api/applications").status_code == 401
    no_auth_client = TestClient(app)
    no_auth_client.headers["Authorization"] = "Bearer garbage-token"
    assert no_auth_client.get("/api/applications").status_code == 401


def test_admin_can_create_user_and_data_is_isolated(client):
    # client is logged in as the bootstrap admin
    create_res = client.post("/api/users", json={"username": "second_user", "password": "secondpassword123"})
    assert create_res.status_code == 200
    assert create_res.json()["username"] == "second_user"
    assert create_res.json()["is_admin"] is False

    # A non-admin cannot create more users
    second_client = _login_client("second_user", "secondpassword123")
    forbidden_res = second_client.post("/api/users", json={"username": "third_user", "password": "whatever123"})
    assert forbidden_res.status_code == 403

    # Admin saves a profile; second user should NOT see it
    client.post(
        "/api/profile",
        json={
            "full_name": "Admin User",
            "email": "admin@example.com",
            "phone": "",
            "location": "",
            "summary": "Admin's profile",
            "skills": ["Python"],
            "experiences": [],
            "projects": [],
            "certifications": [],
            "education": [],
        },
    )
    second_profile = second_client.get("/api/profile").json()
    assert second_profile["full_name"] == ""  # second user's own (empty) profile, not the admin's

    # Second user saves their own distinct profile
    second_client.post(
        "/api/profile",
        json={
            "full_name": "Second User",
            "email": "second@example.com",
            "phone": "",
            "location": "",
            "summary": "",
            "skills": [],
            "experiences": [],
            "projects": [],
            "certifications": [],
            "education": [],
        },
    )
    admin_profile = client.get("/api/profile").json()
    assert admin_profile["full_name"] == "Admin User"  # unaffected by second user's save

    # Applications created by one user are invisible to (and not fetchable by id by) the other
    app_res = client.post(
        "/api/applications",
        json={"company": "AdminCo", "position": "Role", "job_description": "JD"},
    )
    admin_app_id = app_res.json()["id"]
    assert second_client.get("/api/applications").json() == []
    assert second_client.get(f"/api/applications/{admin_app_id}").status_code == 404


def test_profile_empty_get(client):
    response = client.get("/api/profile")
    assert response.status_code == 200
    data = response.json()
    assert data["full_name"] == ""
    assert data["experiences"] == []
    assert data["skills"] == []


def test_profile_save_and_reload(client):
    payload = {
        "full_name": "Jane Doe",
        "email": "jane@example.com",
        "phone": "+1 555 123 4567",
        "location": "New York, NY",
        "linkedin": "https://linkedin.com/in/janedoe",
        "portfolio_url": "https://janedoe.dev",
        "summary": "Experienced Full Stack Engineer with 7 years in cloud systems.",
        "skills": ["Python", "FastAPI", "React", "TypeScript", "PostgreSQL", "Docker"],
        "experiences": [
            {
                "company": "Tech Corp",
                "title": "Lead Software Engineer",
                "start_date": "Jan 2021",
                "end_date": "Present",
                "bullet_points": [
                    "Architected high-throughput microservices using FastAPI and Kafka.",
                    "Improved system uptime from 99.2% to 99.99%.",
                ],
            }
        ],
        "projects": [
            {
                "name": "Cloud Monitor",
                "description": "Real-time infrastructure observability dashboard.",
                "tech_stack": "Go, React, Redis",
                "link": "https://github.com/janedoe/cloud-monitor",
            }
        ],
        "certifications": [
            {
                "name": "AWS Certified Solutions Architect",
                "issuer": "Amazon Web Services",
                "date_earned": "2023",
            }
        ],
        "education": [
            {
                "institution": "Columbia University",
                "degree": "B.S.",
                "field": "Computer Science",
                "graduation_year": "2017",
            }
        ],
    }

    # Save Profile
    save_res = client.post("/api/profile", json=payload)
    assert save_res.status_code == 200
    saved_data = save_res.json()
    assert saved_data["full_name"] == "Jane Doe"
    assert len(saved_data["experiences"]) == 1
    assert saved_data["experiences"][0]["company"] == "Tech Corp"
    assert len(saved_data["skills"]) == 6

    # Reload Profile
    get_res = client.get("/api/profile")
    assert get_res.status_code == 200
    reloaded_data = get_res.json()
    assert reloaded_data["full_name"] == "Jane Doe"
    assert reloaded_data["summary"] == payload["summary"]
    assert len(reloaded_data["experiences"]) == 1
    assert len(reloaded_data["experiences"][0]["bullet_points"]) == 2
    assert len(reloaded_data["projects"]) == 1
    assert len(reloaded_data["certifications"]) == 1
    assert len(reloaded_data["education"]) == 1


def test_settings_save_and_get(client):
    payload = {
        "provider": "openai",
        "api_key": "sk-test-mock-key-12345",
        "model": "gpt-4o-mini",
        "is_default": True,
    }
    post_res = client.post("/api/settings", json=payload)
    assert post_res.status_code == 200
    saved = post_res.json()
    assert saved["provider"] == "openai"
    assert saved["api_key"] == "sk-test-mock-key-12345"
    assert saved["is_default"] is True

    get_res = client.get("/api/settings")
    assert get_res.status_code == 200
    settings = get_res.json()
    assert len(settings) >= 1
    assert settings[0]["provider"] == "openai"


def test_groq_settings_save_and_routing(client):
    # Save Groq settings
    payload = {
        "provider": "groq",
        "api_key": "gsk_mock_test_key_12345",
        "model": "openai/gpt-oss-120b",
        "is_default": True,
    }
    post_res = client.post("/api/settings", json=payload)
    assert post_res.status_code == 200
    saved = post_res.json()
    assert saved["provider"] == "groq"
    assert saved["model"] == "openai/gpt-oss-120b"
    assert saved["is_default"] is True

    db = TestingSessionLocal()
    try:
        resolved = ai_router.get_provider(db, user_id=client.test_user_id)
        assert isinstance(resolved.provider, GroqProvider)
        assert resolved.provider.model == "openai/gpt-oss-120b"
        assert resolved.provider.api_key == "gsk_mock_test_key_12345"
        assert resolved.provider_name == "groq"
        assert resolved.model == "openai/gpt-oss-120b"
    finally:
        db.close()


def _mock_get_provider(db, user_id=None, provider_override=None):
    return ai_router.ResolvedProvider(
        provider=MockAiProvider(), provider_name="groq", model="llama-3.3-70b-versatile"
    )


class MockAiProvider(AiProvider):
    def generate(self, prompt: str, system_prompt: str = "") -> str:
        if "HR data parsing assistant" in system_prompt:
            return json.dumps({
                "full_name": "Jordan Smith",
                "email": "jordan.smith@example.com",
                "phone": "+1 555 333 4444",
                "location": "Austin, TX",
                "linkedin": "https://linkedin.com/in/jordansmith",
                "portfolio_url": "https://jordansmith.dev",
                "summary": "Full Stack Engineer with expertise in modern Python & React architectures.",
                "skills": ["Python", "Django", "FastAPI", "React", "Docker", "PostgreSQL"],
                "experiences": [
                    {
                        "company": "Apex Dynamics",
                        "title": "Senior Engineer",
                        "start_date": "2020",
                        "end_date": "Present",
                        "bullet_points": [
                            "Built scalable backend services handling millions of events daily.",
                            "Mentored 4 junior engineers."
                        ]
                    }
                ],
                "projects": [
                    {
                        "name": "Metrics Engine",
                        "description": "Distributed analytics pipeline.",
                        "tech_stack": "Python, Redis, Docker",
                        "link": "https://github.com/jordansmith/metrics"
                    }
                ],
                "certifications": [
                    {
                        "name": "CKA - Certified Kubernetes Administrator",
                        "issuer": "CNCF",
                        "date_earned": "2022"
                    }
                ],
                "education": [
                    {
                        "institution": "University of Texas at Austin",
                        "degree": "B.S.",
                        "field": "Computer Engineering",
                        "graduation_year": "2019"
                    }
                ]
            })
        elif "STRICT FACTUALITY AND INTEGRITY RULES" in system_prompt or "OUTPUT JSON SCHEMA" in system_prompt:
            return json.dumps({
                "summary": "Tailored executive summary matching target job.",
                "skills": ["Python", "FastAPI", "React"],
                "experience": [
                    {
                        "company": "Tech Corp",
                        "title": "Lead Software Engineer",
                        "start_date": "Jan 2021",
                        "end_date": "Present",
                        "bullet_points": ["Spearheaded microservices optimization matching job needs."]
                    }
                ],
                "projects": [
                    {
                        "name": "Cloud Monitor",
                        "description": "High availability dashboard.",
                        "tech_stack": "Go, React, Redis",
                        "link": "https://github.com/janedoe/cloud-monitor"
                    }
                ],
                "certifications": [],
                "education": []
            })
        else:
            return (
                "Dear Hiring Manager,\n\n"
                "I am excited to bring my experience building scalable distributed systems at Tech Corp to your engineering team.\n\n"
                "In my previous role, I architected high-throughput microservices using FastAPI and Kafka, improving uptime to 99.99%.\n\n"
                "I look forward to discussing how my background aligns with your vision."
            )


def test_generate_endpoint_with_mock_provider(client, monkeypatch):
    profile_payload = {
        "full_name": "Jane Doe",
        "email": "jane@example.com",
        "phone": "+1 555 123 4567",
        "location": "New York, NY",
        "summary": "Experienced Full Stack Engineer with 7 years in cloud systems.",
        "skills": ["Python", "FastAPI", "React"],
        "experiences": [
            {
                "company": "Tech Corp",
                "title": "Lead Software Engineer",
                "start_date": "Jan 2021",
                "end_date": "Present",
                "bullet_points": ["Architected high-throughput microservices."],
            }
        ],
        "projects": [],
        "certifications": [],
        "education": [],
    }
    client.post("/api/profile", json=profile_payload)

    settings_payload = {
        "provider": "groq",
        "api_key": "gsk_test_key",
        "model": "llama-3.3-70b-versatile",
        "is_default": True,
    }
    client.post("/api/settings", json=settings_payload)

    monkeypatch.setattr(ai_router, "get_provider", _mock_get_provider)

    gen_payload = {
        "company": "InnovateTech",
        "position": "Staff Backend Engineer",
        "job_description": "We are seeking a Staff Backend Engineer proficient in Python, FastAPI, distributed systems, and cloud architecture.",
    }
    gen_res = client.post("/api/generate", json=gen_payload)
    assert gen_res.status_code == 200
    result = gen_res.json()
    assert "cv" in result
    assert "cover_letter" in result
    assert result["cv"]["summary"] == "Tailored executive summary matching target job."
    assert "Dear Hiring Manager" in result["cover_letter"]
    assert "Tech Corp" in result["cover_letter"]
    assert result["application_id"] is not None

    # The generation should have created a tracked Application row
    app_res = client.get(f"/api/applications/{result['application_id']}")
    assert app_res.status_code == 200
    app_data = app_res.json()
    assert app_data["company"] == "InnovateTech"
    assert app_data["position"] == "Staff Backend Engineer"
    assert app_data["provider_used"] == "groq"
    assert app_data["model_used"] == "llama-3.3-70b-versatile"
    assert app_data["status"] == "Generated"
    assert app_data["generated_cv"]["summary"] == "Tailored executive summary matching target job."


def test_upload_cv_docx_and_unsupported(client):
    doc = docx.Document()
    doc.add_paragraph("Alex Morgan - Full Stack Developer")
    doc.add_paragraph("Skills: Python, TypeScript, React")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    res = client.post(
        "/api/profile/upload-cv",
        files={"file": ("resume.docx", doc_io.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert res.status_code == 200
    assert "Alex Morgan - Full Stack Developer" in res.json()["text"]

    bad_res = client.post(
        "/api/profile/upload-cv",
        files={"file": ("resume.txt", b"plain text content", "text/plain")},
    )
    assert bad_res.status_code == 400


def test_upload_cv_pdf(client):
    pdf = pdfium.PdfDocument.new()
    pdf.new_page(width=595, height=842)
    buf = io.BytesIO()
    pdf.save(buf)
    buf.seek(0)

    res = client.post(
        "/api/profile/upload-cv",
        files={"file": ("resume.pdf", buf.getvalue(), "application/pdf")},
    )
    assert res.status_code == 200
    assert "text" in res.json()


def test_parse_cv_endpoint(client, monkeypatch):
    # Setup active setting with Groq
    settings_payload = {
        "provider": "groq",
        "api_key": "gsk_test_key",
        "model": "llama-3.3-70b-versatile",
        "is_default": True,
    }
    client.post("/api/settings", json=settings_payload)
    monkeypatch.setattr(ai_router, "get_provider", _mock_get_provider)

    # Create mock docx
    doc = docx.Document()
    doc.add_heading("Jordan Smith", level=1)
    doc.add_paragraph("jordan.smith@example.com | +1 555 333 4444 | Austin, TX")
    doc.add_paragraph("Experience: Senior Engineer at Apex Dynamics (2020 - Present)")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    parse_res = client.post(
        "/api/profile/parse-cv",
        files={"file": ("jordan_resume.docx", doc_io.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert parse_res.status_code == 200
    data = parse_res.json()
    assert "parsed_profile" in data
    assert "raw_text" in data
    parsed = data["parsed_profile"]
    assert parsed["full_name"] == "Jordan Smith"
    assert parsed["email"] == "jordan.smith@example.com"
    assert len(parsed["skills"]) == 6
    assert len(parsed["experiences"]) == 1
    assert parsed["experiences"][0]["company"] == "Apex Dynamics"
    assert len(parsed["projects"]) == 1
    assert len(parsed["certifications"]) == 1
    assert len(parsed["education"]) == 1


def test_applications_crud(client):
    create_payload = {
        "company": "Globex",
        "position": "Backend Engineer",
        "job_description": "Build APIs.",
        "provider_used": "openai",
        "model_used": "gpt-4o-mini",
        "generated_cv": {"summary": "A summary", "skills": ["Python"]},
        "generated_cover_letter": "Dear Hiring Manager, ...",
    }
    create_res = client.post("/api/applications", json=create_payload)
    assert create_res.status_code == 200
    created = create_res.json()
    assert created["company"] == "Globex"
    assert created["status"] == "Generated"
    assert created["generated_cv"]["summary"] == "A summary"
    app_id = created["id"]

    # List, most recent first
    second_res = client.post(
        "/api/applications",
        json={**create_payload, "company": "Initech", "generated_cv": {}},
    )
    assert second_res.status_code == 200
    list_res = client.get("/api/applications")
    assert list_res.status_code == 200
    listing = list_res.json()
    assert len(listing) == 2
    assert listing[0]["company"] == "Initech"  # most recently created first

    # Get single
    get_res = client.get(f"/api/applications/{app_id}")
    assert get_res.status_code == 200
    assert get_res.json()["company"] == "Globex"

    # Update status
    put_res = client.put(f"/api/applications/{app_id}", json={"status": "Interview"})
    assert put_res.status_code == 200
    assert put_res.json()["status"] == "Interview"

    # Invalid status rejected by schema validation
    bad_put_res = client.put(f"/api/applications/{app_id}", json={"status": "NotAStatus"})
    assert bad_put_res.status_code == 422

    # Delete
    delete_res = client.delete(f"/api/applications/{app_id}")
    assert delete_res.status_code == 200
    missing_res = client.get(f"/api/applications/{app_id}")
    assert missing_res.status_code == 404


def test_generate_with_no_default_provider_but_override(client, monkeypatch):
    profile_payload = {
        "full_name": "Jane Doe",
        "email": "jane@example.com",
        "phone": "",
        "location": "",
        "summary": "Summary",
        "skills": ["Python"],
        "experiences": [],
        "projects": [],
        "certifications": [],
        "education": [],
    }
    client.post("/api/profile", json=profile_payload)

    # Save two providers, only "openai" marked default
    client.post(
        "/api/settings",
        json={"provider": "openai", "api_key": "sk-a", "model": "gpt-4o-mini", "is_default": True},
    )
    client.post(
        "/api/settings",
        json={"provider": "groq", "api_key": "gsk-b", "model": "openai/gpt-oss-120b", "is_default": False},
    )

    captured = {}

    def fake_get_provider(db, user_id=None, provider_override=None):
        captured["override"] = provider_override
        return ai_router.ResolvedProvider(
            provider=MockAiProvider(), provider_name=provider_override or "openai", model="test-model"
        )

    monkeypatch.setattr(ai_router, "get_provider", fake_get_provider)

    gen_payload = {
        "company": "Acme",
        "position": "Engineer",
        "job_description": "Build things with Python and FastAPI.",
        "provider": "groq",
    }
    gen_res = client.post("/api/generate", json=gen_payload)
    assert gen_res.status_code == 200
    assert captured["override"] == "groq"
    assert gen_res.json()["application_id"] is not None
