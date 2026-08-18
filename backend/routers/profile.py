import io
import json
import re
from typing import List, Any, Dict
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
import pdfplumber
import docx

from database import get_db
import models
import schemas
import services.ai.router as ai_router
from services.prompts.parse_cv import build_parse_cv_prompt
from auth import require_app_password

router = APIRouter(prefix="/api/profile", tags=["Profile"], dependencies=[Depends(require_app_password)])


def _as_list(value: Any) -> List[Any]:
    """AI output sometimes returns null instead of [] for an empty section; normalize it."""
    return value if isinstance(value, list) else []


def _clean_json_output(text: str) -> str:
    text = text.strip()
    match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", text, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return text


def _serialize_profile(profile: models.Profile) -> Dict[str, Any]:
    skills_list = []
    if profile.skills:
        try:
            skills_list = json.loads(profile.skills)
            if not isinstance(skills_list, list):
                skills_list = [str(profile.skills)]
        except Exception:
            skills_list = [s.strip() for s in profile.skills.split(",") if s.strip()]

    experiences = []
    for exp in profile.experiences:
        bullets = []
        if exp.bullet_points:
            try:
                bullets = json.loads(exp.bullet_points)
                if not isinstance(bullets, list):
                    bullets = [str(exp.bullet_points)]
            except Exception:
                bullets = [b.strip() for b in exp.bullet_points.split("\n") if b.strip()]
        experiences.append(
            {
                "id": exp.id,
                "company": exp.company or "",
                "title": exp.title or "",
                "start_date": exp.start_date or "",
                "end_date": exp.end_date,
                "bullet_points": bullets,
            }
        )

    projects = [
        {
            "id": p.id,
            "name": p.name or "",
            "description": p.description or "",
            "tech_stack": p.tech_stack or "",
            "link": p.link,
        }
        for p in profile.projects
    ]

    certifications = [
        {
            "id": c.id,
            "name": c.name or "",
            "issuer": c.issuer or "",
            "date_earned": c.date_earned,
        }
        for c in profile.certifications
    ]

    education = [
        {
            "id": e.id,
            "institution": e.institution or "",
            "degree": e.degree or "",
            "field": e.field,
            "graduation_year": e.graduation_year,
        }
        for e in profile.education
    ]

    return {
        "id": profile.id,
        "full_name": profile.full_name or "",
        "email": profile.email or "",
        "phone": profile.phone or "",
        "location": profile.location or "",
        "linkedin": profile.linkedin or "",
        "portfolio_url": profile.portfolio_url or "",
        "summary": profile.summary or "",
        "skills": skills_list,
        "experiences": experiences,
        "projects": projects,
        "certifications": certifications,
        "education": education,
    }


def _extract_text_from_file(filename: str, contents: bytes) -> str:
    extracted_text = ""
    filename = filename.lower()

    if filename.endswith(".pdf"):
        try:
            with pdfplumber.open(io.BytesIO(contents)) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                extracted_text = "\n\n".join(pages_text)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")
    elif filename.endswith(".docx"):
        try:
            doc = docx.Document(io.BytesIO(contents))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
            extracted_text = "\n\n".join(paragraphs)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Please upload a PDF (.pdf) or Word document (.docx).",
        )

    if not extracted_text.strip():
        extracted_text = "(No readable text could be extracted from this document.)"

    return extracted_text.strip()


@router.get("", response_model=schemas.ProfileSchema)
def get_profile(db: Session = Depends(get_db)):
    profile = db.query(models.Profile).first()
    if not profile:
        return schemas.ProfileSchema(
            id=None,
            full_name="",
            email="",
            phone="",
            location="",
            linkedin="",
            portfolio_url="",
            summary="",
            skills=[],
            experiences=[],
            projects=[],
            certifications=[],
            education=[],
        )
    return _serialize_profile(profile)


@router.post("", response_model=schemas.ProfileSchema)
@router.put("", response_model=schemas.ProfileSchema)
def save_profile(data: schemas.ProfileSaveRequest, db: Session = Depends(get_db)):
    profile = db.query(models.Profile).first()
    if not profile:
        profile = models.Profile()
        db.add(profile)
        db.flush()

    # Update basic profile fields
    profile.full_name = data.full_name
    profile.email = data.email
    profile.phone = data.phone
    profile.location = data.location
    profile.linkedin = data.linkedin
    profile.portfolio_url = data.portfolio_url
    profile.summary = data.summary
    profile.skills = json.dumps(data.skills)

    # Clear existing child records
    db.query(models.Experience).filter(models.Experience.profile_id == profile.id).delete()
    db.query(models.Project).filter(models.Project.profile_id == profile.id).delete()
    db.query(models.Certification).filter(models.Certification.profile_id == profile.id).delete()
    db.query(models.Education).filter(models.Education.profile_id == profile.id).delete()

    # Add updated experiences
    for exp_data in data.experiences:
        exp = models.Experience(
            profile_id=profile.id,
            company=exp_data.company,
            title=exp_data.title,
            start_date=exp_data.start_date,
            end_date=exp_data.end_date,
            bullet_points=json.dumps(exp_data.bullet_points),
        )
        db.add(exp)

    # Add updated projects
    for proj_data in data.projects:
        proj = models.Project(
            profile_id=profile.id,
            name=proj_data.name,
            description=proj_data.description,
            tech_stack=proj_data.tech_stack,
            link=proj_data.link,
        )
        db.add(proj)

    # Add updated certifications
    for cert_data in data.certifications:
        cert = models.Certification(
            profile_id=profile.id,
            name=cert_data.name,
            issuer=cert_data.issuer,
            date_earned=cert_data.date_earned,
        )
        db.add(cert)

    # Add updated education
    for edu_data in data.education:
        edu = models.Education(
            profile_id=profile.id,
            institution=edu_data.institution,
            degree=edu_data.degree,
            field=edu_data.field,
            graduation_year=edu_data.graduation_year,
        )
        db.add(edu)

    db.commit()
    db.refresh(profile)
    return _serialize_profile(profile)


@router.post("/upload-cv", response_model=schemas.UploadCvResponse)
async def upload_cv(file: UploadFile = File(...)):
    filename = file.filename or ""
    contents = await file.read()
    extracted_text = _extract_text_from_file(filename, contents)
    return {"text": extracted_text}


@router.post("/parse-cv", response_model=schemas.ParseCvResponse)
async def parse_cv_with_ai(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    filename = file.filename or ""
    contents = await file.read()
    extracted_text = _extract_text_from_file(filename, contents)

    if not extracted_text or extracted_text.startswith("(No readable text"):
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from the uploaded document.",
        )

    # Load AI Provider
    resolved = ai_router.get_provider(db)
    ai_provider = resolved.provider

    # Build parse prompt
    system_prompt, user_prompt = build_parse_cv_prompt(extracted_text)

    try:
        raw_response = ai_provider.generate(prompt=user_prompt, system_prompt=system_prompt)
    except Exception as e:
        raise HTTPException(
            status_code=502,
            detail=f"AI Provider error during resume parsing: {str(e)}",
        )

    cleaned_json = _clean_json_output(raw_response)

    try:
        parsed_data = json.loads(cleaned_json)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse AI output into valid profile structure: {str(e)}",
        )

    if not isinstance(parsed_data, dict):
        raise HTTPException(
            status_code=500,
            detail="AI output was not a valid profile object. Please try again.",
        )

    # Format into ProfileSaveRequest schema structure with defaults
    parsed_profile = schemas.ProfileSaveRequest(
        full_name=parsed_data.get("full_name", "") or "",
        email=parsed_data.get("email", "") or "",
        phone=parsed_data.get("phone", "") or "",
        location=parsed_data.get("location", "") or "",
        linkedin=parsed_data.get("linkedin"),
        portfolio_url=parsed_data.get("portfolio_url"),
        summary=parsed_data.get("summary", "") or "",
        skills=[s for s in _as_list(parsed_data.get("skills")) if isinstance(s, str) and s.strip()],
        experiences=[
            schemas.ExperienceBase(
                company=exp.get("company", "") or "",
                title=exp.get("title", "") or "",
                start_date=exp.get("start_date", "") or "",
                end_date=exp.get("end_date"),
                bullet_points=[b for b in _as_list(exp.get("bullet_points")) if isinstance(b, str) and b.strip()],
            )
            for exp in _as_list(parsed_data.get("experiences"))
            if isinstance(exp, dict)
        ],
        projects=[
            schemas.ProjectBase(
                name=proj.get("name", "") or "",
                description=proj.get("description", "") or "",
                tech_stack=proj.get("tech_stack", "") or "",
                link=proj.get("link"),
            )
            for proj in _as_list(parsed_data.get("projects"))
            if isinstance(proj, dict)
        ],
        certifications=[
            schemas.CertificationBase(
                name=cert.get("name", "") or "",
                issuer=cert.get("issuer", "") or "",
                date_earned=cert.get("date_earned"),
            )
            for cert in _as_list(parsed_data.get("certifications"))
            if isinstance(cert, dict)
        ],
        education=[
            schemas.EducationBase(
                institution=edu.get("institution", "") or "",
                degree=edu.get("degree", "") or "",
                field=edu.get("field"),
                graduation_year=edu.get("graduation_year"),
            )
            for edu in _as_list(parsed_data.get("education"))
            if isinstance(edu, dict)
        ],
    )

    return schemas.ParseCvResponse(
        parsed_profile=parsed_profile,
        raw_text=extracted_text,
    )
